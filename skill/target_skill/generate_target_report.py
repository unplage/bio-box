#!/usr/bin/env python3
"""
生物药靶点专利调研报告生成器 - 通用模板
用法：python3 generate_target_report.py --target PD-1 --time "2021-2026"
"""

import argparse
import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_heading_styled(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
    return heading

def add_table_styled(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, '003366')
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9)
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
            if row_idx % 2 == 0:
                set_cell_shading(cell, 'E8F0FE')
    doc.add_paragraph('')
    return table

def add_evidence(doc, claim, source_name, url):
    """添加带证据的要点"""
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(claim)
    run.font.size = Pt(10)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(1.5)
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(6)
    run2 = p2.add_run(f'来源：{source_name}')
    run2.font.size = Pt(8)
    run2.font.color.rgb = RGBColor(0, 102, 204)
    run2.font.italic = True
    run3 = p2.add_run(f'\n{url}')
    run3.font.size = Pt(7)
    run3.font.color.rgb = RGBColor(0, 102, 204)

def create_report(target_name, time_range="2021-2026", region="全球"):
    """创建靶点专利调研报告"""
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    
    # 封面
    for _ in range(6):
        doc.add_paragraph('')
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f'{target_name}靶点专利调研报告')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f'（{time_range}年全球专利分析）')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(102, 102, 102)
    doc.add_paragraph('')
    doc.add_paragraph('')
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info_para.add_run(f'报告日期：{datetime.date.today().strftime("%Y年%m月%d日")}')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(102, 102, 102)
    doc.add_paragraph('')
    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disclaimer.add_run('【机密文件】仅供内部研究使用')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(255, 0, 0)
    run.font.bold = True
    doc.add_page_break()
    
    # 目录
    doc.add_heading('目  录', level=1)
    toc = [
        f'一、{target_name}技术概述',
        '二、专利检索策略',
        '三、全球专利申请趋势',
        '四、专利区域分布',
        '五、主要申请人排名',
        '六、重点竞争对手分析',
        '七、技术路线分布',
        '八、核心技术专利识别',
        '九、结论与建议（含引用证据）',
        '附录：参考文献'
    ]
    for item in toc:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.bold = True
    doc.add_page_break()
    
    # 一、技术概述
    add_heading_styled(doc, f'一、{target_name}技术概述', level=1)
    add_heading_styled(doc, f'1.1 {target_name}靶点介绍', level=2)
    doc.add_paragraph(
        f'{target_name}靶点是肿瘤免疫治疗/自身免疫病治疗领域的重要靶点。'
        f'本报告将对{target_name}靶点的全球专利布局、竞争格局和技术发展趋势进行系统分析。'
    )
    
    add_heading_styled(doc, '1.2 治疗机制与临床应用', level=2)
    doc.add_paragraph(f'目前针对{target_name}靶点的治疗策略主要包括以下几种形式：')
    
    # TODO: 根据具体靶点填写治疗模式
    headers = ['技术类型', '代表药物', '作用机制', '获批适应症']
    rows = [
        ['[待填写]', '[待填写]', '[待填写]', '[待填写]'],
    ]
    add_table_styled(doc, headers, rows)
    
    doc.add_page_break()
    
    # 二、专利检索策略
    add_heading_styled(doc, '二、专利检索策略', level=1)
    add_heading_styled(doc, '2.1 检索关键词', level=2)
    doc.add_paragraph('本报告采用以下关键词组合进行专利检索：')
    
    keywords = [
        f'英文关键词：{target_name}, {target_name} antibody, CAR-T, chimeric antigen receptor, bispecific antibody, ADC',
        f'中文关键词：{target_name}, 抗体, 嵌合抗原受体, 双特异性抗体, 抗体偶联药物',
        f'药物名称：[待填写已上市药物名称]',
    ]
    for kw in keywords:
        doc.add_paragraph(kw, style='List Bullet')
    
    add_heading_styled(doc, '2.2 检索数据库', level=2)
    doc.add_paragraph(
        '本报告数据来源包括：Google Patents、WIPO PatentScope、USPTO、EPO、CNIPA、'
        'Derwent Innovation、Lens.org等公开专利数据库，以及Patsnap、Eureka等专利分析平台。'
    )
    
    doc.add_page_break()
    
    # 三、全球专利申请趋势
    add_heading_styled(doc, '三、全球专利申请趋势', level=1)
    
    # TODO: 根据实际数据填写
    headers = ['年份', '全球申请量（估计）', '同比增长', '主要技术方向']
    rows = [
        ['[年份]', '[数量]', '[增长率]', '[技术方向]'],
    ]
    add_table_styled(doc, headers, rows)
    
    doc.add_page_break()
    
    # 四、专利区域分布
    add_heading_styled(doc, '四、专利区域分布', level=1)
    
    headers = ['国家/地区', '专利申请占比', '主要申请人类型', '技术特点']
    rows = [
        ['[国家]', '[占比]', '[类型]', '[特点]'],
    ]
    add_table_styled(doc, headers, rows)
    
    doc.add_page_break()
    
    # 五、主要申请人排名
    add_heading_styled(doc, '五、主要申请人排名', level=1)
    
    headers = ['排名', '申请人', '国家', '专利数量(估)', '主要技术方向']
    rows = [
        ['[排名]', '[申请人]', '[国家]', '[数量]', '[技术方向]'],
    ]
    add_table_styled(doc, headers, rows)
    
    doc.add_page_break()
    
    # 六、重点竞争对手分析
    add_heading_styled(doc, '六、重点竞争对手分析', level=1)
    
    # TODO: 根据实际调研结果填写
    competitors = ['[企业1]', '[企业2]', '[企业3]']
    for i, competitor in enumerate(competitors, 1):
        add_heading_styled(doc, f'6.{i} {competitor}', level=2)
        doc.add_paragraph(f'{competitor}是{target_name}领域的重要参与者。')
        headers = ['分析维度', '详细信息']
        rows = [
            ['获批产品', '[待填写]'],
            ['技术路线', '[待填写]'],
            ['适应症', '[待填写]'],
            ['专利数量', '[待填写]'],
        ]
        add_table_styled(doc, headers, rows)
    
    doc.add_page_break()
    
    # 七、技术路线分布
    add_heading_styled(doc, '七、技术路线分布', level=1)
    
    # TODO: 根据实际调研结果填写
    tech_routes = ['[技术路线1]', '[技术路线2]', '[技术路线3]']
    for i, route in enumerate(tech_routes, 1):
        add_heading_styled(doc, f'7.{i} {route}', level=2)
        doc.add_paragraph(f'{route}是{target_name}领域的重要技术方向。')
    
    doc.add_page_break()
    
    # 八、核心技术专利识别
    add_heading_styled(doc, '八、核心技术专利识别', level=1)
    
    headers = ['专利号', '标题', '权利人', '优先权日', '技术要点']
    rows = [
        ['[专利号]', '[标题]', '[权利人]', '[日期]', '[要点]'],
    ]
    add_table_styled(doc, headers, rows)
    
    doc.add_page_break()
    
    # 九、结论与建议
    add_heading_styled(doc, '九、结论与建议（含引用证据）', level=1)
    
    add_heading_styled(doc, '9.1 竞争格局总结', level=2)
    
    # TODO: 根据实际调研结果填写结论和证据
    p = doc.add_paragraph()
    run = p.add_run(f'结论1：{target_name}领域呈现[待填写]的竞争格局')
    run.font.bold = True
    run.font.size = Pt(11)
    
    add_evidence(doc,
        '[待填写结论支撑数据]',
        '[来源名称]',
        '[来源URL]')
    
    doc.add_page_break()
    
    # 附录：参考文献
    add_heading_styled(doc, '附录：参考文献', level=1)
    
    # TODO: 根据实际引用来源填写
    refs = [
        '1. [来源1]',
        '2. [来源2]',
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.font.size = Pt(9)
    
    return doc

def main():
    parser = argparse.ArgumentParser(description='生成靶点专利调研报告')
    parser.add_argument('--target', required=True, help='靶点名称，如PD-1、HER2等')
    parser.add_argument('--time', default='2021-2026', help='时间范围')
    parser.add_argument('--region', default='全球', help='地域范围')
    parser.add_argument('--output', help='输出文件路径')
    
    args = parser.parse_args()
    
    doc = create_report(args.target, args.time, args.region)
    
    if args.output:
        output_path = args.output
    else:
        output_path = f'{args.target}靶点专利调研报告_{args.time}_模板.docx'
    
    doc.save(output_path)
    print(f'报告已生成：{output_path}')

if __name__ == '__main__':
    main()